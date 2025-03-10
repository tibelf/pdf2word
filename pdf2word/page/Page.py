# -*- coding: utf-8 -*-

'''Page object parsed with PDF raw dict.

In addition to base structure described in :py:class:`~pdf2docx.page.RawPage`, 
some new features, e.g. sections, table block, are also included. 
Page elements structure:

* :py:class:`~pdf2docx.page.Page` >> :py:class:`~pdf2docx.layout.Section` >> :py:class:`~pdf2docx.layout.Column`  
    * :py:class:`~pdf2docx.layout.Blocks`
        * :py:class:`~pdf2docx.text.TextBlock` >> 
          :py:class:`~pdf2docx.text.Line` >> 
          :py:class:`~pdf2docx.text.TextSpan` / :py:class:`~pdf2docx.image.ImageSpan` >>
          :py:class:`~pdf2docx.text.Char`
        * :py:class:`~pdf2docx.table.TableBlock` >>
          :py:class:`~pdf2docx.table.Row` >> 
          :py:class:`~pdf2docx.table.Cell`
            * :py:class:`~pdf2docx.layout.Blocks`
            * :py:class:`~pdf2docx.shape.Shapes`
    * :py:class:`~pdf2docx.shape.Shapes`
        * :py:class:`~pdf2docx.shape.Shape.Stroke`
        * :py:class:`~pdf2docx.shape.Shape.Fill`
        * :py:class:`~pdf2docx.shape.Shape.Hyperlink`

::

    {
        "id": 0, # page index
        "width" : w,
        "height": h,
        "margin": [left, right, top, bottom],
        "sections": [{
            ... # section properties
        }, ...],
        "floats": [{
            ... # floating picture
        }, ...],
        "formulas": [{
            "bbox": [x0, y0, x1, y1],
            "latex": "..."
        }, ...]
    }

'''

from docx.shared import Pt
from docx.enum.section import WD_SECTION
from ..common.Collection import BaseCollection
from ..common.share import debug_plot
from .BasePage import BasePage
from ..layout.Sections import Sections
from ..image.ImageBlock import ImageBlock
import logging


class Page(BasePage):
    '''Object representing the whole page, e.g. margins, sections.'''

    def __init__(self, id:int=-1, 
                        skip_parsing:bool=True,
                        width:float=0.0,
                        height:float=0.0,
                        header:str=None, 
                        footer:str=None, 
                        margin:tuple=None, 
                        sections:Sections=None,
                        float_images:BaseCollection=None,
                        formulas:list=None):
        '''Initialize page layout.

        Args:
            id (int, optional): Page index. Defaults to -1.
            skip_parsing (bool, optional): Don't parse page if True. Defaults to True.
            width (float, optional): Page width. Defaults to 0.0.
            height (float, optional): Page height. Defaults to 0.0.
            header (str, optional): Page header. Defaults to None.
            footer (str, optional): Page footer. Defaults to None.
            margin (tuple, optional): Page margin. Defaults to None.
            sections (Sections, optional): Page contents. Defaults to None.
            float_images (BaseCollection, optional): Float images in th is page. Defaults to None.
            formulas (list, optional): Mathematical formulas in this page. Defaults to None.
        ''' 
        # page index
        self.id = id
        self.skip_parsing = skip_parsing

        # page size and margin
        super().__init__(width=width, height=height, margin=margin)

        # flow structure: 
        # Section -> Column -> Blocks -> TextBlock/TableBlock
        # TableBlock -> Row -> Cell -> Blocks
        self.sections = sections or Sections(parent=self)

        # page header, footer
        self.header = header or ''
        self.footer = footer or ''
        
        # floating images are separate node under page
        self.float_images = float_images or BaseCollection()
        
        # mathematical formulas
        self.formulas = formulas or []

        self._finalized = False


    @property
    def finalized(self): return self._finalized   


    def store(self):
        '''Store parsed layout data.'''
        res = {
            'id'     : self.id,
            'width'  : self.width,
            'height' : self.height,
            'margin' : self.margin,
            'header' : self.header,
            'footer' : self.footer
        }

        # sections
        res.update({
            'sections': self.sections.store()
        })

        # float images
        if self.float_images:
            res.update({
                'float_images': self.float_images.store()
            })
        
        # formulas
        if self.formulas:
            res.update({
                'formulas': self.formulas
            })

        return res


    def restore(self, data:dict):
        '''Restore Layout from parsed results.'''
        # page id
        self.id = data.get('id', -1)

        # page width/height
        self.width = data.get('width', 0.0)
        self.height = data.get('height', 0.0)

        # page margin
        self.margin = data.get('margin', (0,0,0,0))

        # page header, footer
        self.header = data.get('header', '')
        self.footer = data.get('footer', '')

        # sections
        self.sections.restore(data.get('sections', []))

        # float images
        if 'float_images' in data:
            self._restore_float_images(data.get('float_images', []))
        
        # formulas
        self.formulas = data.get('formulas', [])

        # mark page as finalized
        self._finalized = True


    @debug_plot('Final Layout')
    def parse(self, **settings):
        '''Parse page layout.'''
        # Raised when page layout is invalid or nothing extracted
        if self.skip_parsing:
            return False 

        # already finalized
        if self._finalized: return True

        try:
            # Parse the page content using raw_dict
            raw_dict = settings.get('raw_dict')
            if not raw_dict:
                logging.error('No raw dict provided for parsing page %d', self.id)
                return False

            # parse sections
            status = self.sections.parse(**settings)
            
            # Always mark as finalized for now, even if status is False
            # This ensures the page is included in the output
            self._finalized = True
            
            # process formulas if enabled
            if settings.get('process_formulas', False):
                self._process_formulas(**settings)

            return True
            
        except Exception as e:
            logging.error('Error parsing page %d: %s', self.id, e)
            return False


    def extract_tables(self, **settings):
        '''Extract table contents from page. 
        
        Args:
            settings (dict): Parsing parameters.

        Returns:
            list: A list of ``tabula.Table`` instances.
        '''
        # initialize empty tables
        tables = []

        # collect all table blocks
        collections = []
        for section in self.sections:
            for column in section:
                if settings['extract_stream_table']:
                    collections.extend(column.blocks.table_blocks)
                else:
                    collections.extend(column.blocks.lattice_table_blocks)
        
        # check table
        tables = [] # type: list[ list[list[str]] ]
        for table_block in collections:
            tables.append(table_block.text)

        return tables


    def make_docx(self, doc):
        '''Set page size, margin, and create page. 

        .. note::
            Before running this method, the page layout must be either parsed from source 
            page or restored from parsed data.
        
        Args:
            doc (Document): ``python-docx`` document object
        '''
        # new page
        if doc.paragraphs:
            section = doc.add_section(WD_SECTION.NEW_PAGE)
        else:
            section = doc.sections[0] # a default section is there when opening docx

        # page size
        section.page_width  = Pt(self.width)
        section.page_height = Pt(self.height)

        # page margin
        left,right,top,bottom = self.margin
        section.left_margin = Pt(left)
        section.right_margin = Pt(right)
        section.top_margin = Pt(top)
        section.bottom_margin = Pt(bottom)

        # create flow layout: sections
        self.sections.make_docx(doc)
        
        # add formulas if available
        self._add_formulas_to_docx(doc)

 
    def _restore_float_images(self, raws:list):
        '''Restore float images.'''
        self.float_images.reset()
        for raw in raws:
            image = ImageBlock()
            image.restore(raw)
            image.set_float_image_block()
            self.float_images.append(image)
            
    def _process_formulas(self, **settings):
        '''Detect and process mathematical formulas in the page.'''
        # Only process if formula processing is enabled and formula_processor is available
        if not settings.get('process_formulas', False):
            return
            
        try:
            # Import formula processor
            from ..formula.formula_processor import FormulaProcessor
            
            # Create processor if not exists
            if not hasattr(self, '_formula_processor'):
                self._formula_processor = FormulaProcessor()
                
            # Skip if processor is not enabled
            if not self._formula_processor.enabled:
                return
                
            # Get source page from fitz
            source_page = settings.get('source_page')
            if not source_page:
                return
                
            # Process page to get formulas
            formulas = self._formula_processor.process_page(source_page)
            
            # Store formulas
            self.formulas = formulas
            
        except Exception as e:
            logging.error(f"Error processing formulas on page {self.id}: {e}")
            
    def _add_formulas_to_docx(self, doc):
        '''Add formulas to the docx document.'''
        if not self.formulas:
            return
            
        # Import for creating paragraphs
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        
        # Add each formula as a separate paragraph
        for formula in self.formulas:
            latex = formula.get('latex', '')
            if not latex:
                continue
                
            # Check if it's a display formula (centered) or inline formula
            is_display = latex.startswith('\\[') or latex.startswith('\\begin{')
            
            # Create paragraph for the formula
            p = doc.add_paragraph()
            
            # Center align display formulas
            if is_display:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                
            # Add LaTeX code
            p.add_run(latex)
